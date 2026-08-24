#!/usr/bin/env python3
"""逐字盘点招标文件正文、表格单元格和嵌套表格中的高风险词。"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
from collections import defaultdict
from pathlib import Path

from docx import Document


MARKERS = (
    "▲",
    "●",
    "★",
    "知旅云",
    "课堂资源",
    "云课堂",
    "软件",
    "平台",
    "系统",
    "品牌",
    "厂家",
    "原厂",
    "制造商",
    "开发者",
    "版权",
    "LOGO",
    "授权",
    "核心产品",
    "现场搭建",
    "公示期",
    "终身免费",
    "云服务器",
)

QUOTED_RE = re.compile(r"[“‘‘\"「『]([^”’'\"」』]{2,80})[”’'\"」』]")
NAMED_RE = re.compile(
    r"[\u4e00-\u9fffA-Za-z0-9·&（）()/_-]{2,48}"
    r"(?:课堂资源|云课堂|软件|平台|系统|实训室|资源库|云平台|管理系统|还原卡|应用|模块)"
)


def paragraph_text(paragraph) -> str:
    """读取段落原文；不清理空格、标点、数字或英文大小写。"""
    return paragraph.text


def add_block(blocks: list[dict], location: str, text: str, kind: str) -> None:
    # 保留空段落，确保清单可与文档段落序号逐项对照。
    blocks.append({"location": location, "kind": kind, "text": text})


def walk_table(table, location: str, blocks: list[dict]) -> None:
    """递归遍历单元格和嵌套表格，保留可回到招标文件的行列位置。"""
    for row_index, row in enumerate(table.rows, 1):
        for col_index, cell in enumerate(row.cells, 1):
            cell_location = f"{location}R{row_index}C{col_index}"
            for para_index, paragraph in enumerate(cell.paragraphs, 1):
                add_block(
                    blocks,
                    f"{cell_location}P{para_index}",
                    paragraph_text(paragraph),
                    "table_cell_paragraph",
                )
            for nested_index, nested in enumerate(cell.tables, 1):
                walk_table(nested, f"{cell_location}T{nested_index}", blocks)


def find_converter(name: str) -> str | None:
    """查找文档转换器，优先使用 PATH 中的可执行文件。"""
    return shutil.which(name)


def convert_legacy_doc(path: Path, output_dir: Path) -> tuple[Path, str]:
    """将 .doc 转为 .docx；优先使用 LibreOffice 以保留表格和单元格位置。"""
    soffice = find_converter("soffice")
    if soffice:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", str(output_dir), str(path)],
            capture_output=True,
            text=True,
            check=False,
        )
        converted = output_dir / f"{path.stem}.docx"
        if result.returncode == 0 and converted.exists():
            return converted, "soffice"

    textutil = find_converter("textutil")
    if textutil:
        converted = output_dir / f"{path.stem}.docx"
        result = subprocess.run(
            [textutil, "-convert", "docx", "-output", str(converted), str(path)],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode == 0 and converted.exists():
            return converted, "textutil"

    raise RuntimeError("未找到可用的 .doc 转换器；请安装 LibreOffice 或在 macOS 上提供 textutil。")


def collect_blocks(path: Path) -> tuple[list[dict], dict]:
    conversion = {"source_format": path.suffix.lower(), "method": "none"}
    with tempfile.TemporaryDirectory(prefix="procurement-tender-") as temp_dir:
        source_path = path
        if path.suffix.lower() == ".doc":
            source_path, method = convert_legacy_doc(path, Path(temp_dir))
            conversion["method"] = method
        doc = Document(source_path)
        blocks: list[dict] = []
        for para_index, paragraph in enumerate(doc.paragraphs, 1):
            add_block(blocks, f"P{para_index}", paragraph_text(paragraph), "paragraph")
        for table_index, table in enumerate(doc.tables, 1):
            walk_table(table, f"T{table_index}", blocks)
    return blocks, conversion


def candidate_terms(text: str) -> set[str]:
    """提取需人工复核的专有词候选；完整原文仍保存在 blocks 中。"""
    terms = set(QUOTED_RE.findall(text))
    terms.update(NAMED_RE.findall(text))
    for marker in MARKERS:
        if marker in text:
            terms.add(marker)
    return {term.strip() for term in terms if term.strip()}


def build_inventory(path: Path) -> dict:
    blocks, conversion = collect_blocks(path)
    occurrences: dict[str, list[dict]] = defaultdict(list)
    for block in blocks:
        for term in candidate_terms(block["text"]):
            occurrences[term].append(
                {"location": block["location"], "kind": block["kind"], "context": block["text"]}
            )
    ranked = []
    for term, hits in occurrences.items():
        ranked.append(
            {
                "term": term,
                "occurrences": len(hits),
                "locations": [hit["location"] for hit in hits],
                "contexts": hits,
            }
        )
    ranked.sort(key=lambda item: (-item["occurrences"], item["term"]))
    return {
        "input": str(path),
        "scan_complete": True,
        "conversion": conversion,
        "stats": {
            "paragraphs": sum(block["kind"] == "paragraph" for block in blocks),
            "table_cell_paragraphs": sum(
                block["kind"] == "table_cell_paragraph" for block in blocks
            ),
            "blocks": len(blocks),
            "candidate_terms": len(ranked),
        },
        "candidate_terms": ranked,
        "blocks": blocks,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="逐字扫描DOCX招标文件中的段落、表格和专有词")
    parser.add_argument("input", type=Path, help="招标文件 .docx")
    parser.add_argument("--out", type=Path, help="JSON盘点结果输出路径；省略则打印到标准输出")
    args = parser.parse_args()
    inventory = build_inventory(args.input)
    payload = json.dumps(inventory, ensure_ascii=False, indent=2)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(payload + "\n", encoding="utf-8")
        print(args.out)
    else:
        print(payload)


if __name__ == "__main__":
    main()
